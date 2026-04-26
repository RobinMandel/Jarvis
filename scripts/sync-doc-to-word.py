#!/usr/bin/env python3
"""
Sync-Word-from-GDoc-Content.

Nimmt Plain-Text-Inhalt eines Google Docs (wie er von MCP read_file_content
zurueckkommt) und generiert ein .docx mit unserem Standard-Layout:
  - Erste 6 Zeilen = Absender (rechts-aligned)
  - "Universität Ulm" Block = Empfaenger (links)
  - "Ulm, DD.MM.YYYY" = Datum (rechts)
  - Erste fett-Block = Betreff (bold)
  - Section-Header werden heuristisch erkannt (kurze Zeile, < 80 char, kein
    Satzpunkt am Ende, gefolgt von Doppel-Leerzeile) -> bold
  - Rest = Body-Paragraphs

Aufruf:
    python sync-doc-to-word.py <input-text-file> <output-docx-path>

Plus: aktualisiert `Mission-Control/data/documents.json` mit last_synced.
"""
from __future__ import annotations

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _is_section_header(line: str, next_blank: bool) -> bool:
    """Heuristik: kurze Zeile ohne Satzpunkt am Ende, gefolgt von Leerzeile."""
    t = line.strip()
    if not t or len(t) > 90:
        return False
    if t.endswith((".", ":", ",", "!", "?")):
        return False
    if not next_blank:
        return False
    # Keine Anrede ("Sehr geehrte ...")
    if t.startswith(("Sehr geehrte", "Mit freundlichen")):
        return False
    return True


def _strip_md_escapes(text: str) -> str:
    """`\\+49` -> `+49`, `\\-` -> `-`. Markdown-Escapes aus MCP read entfernen."""
    return re.sub(r"\\([+\-_*\[\]\(\)])", r"\1", text)


def build_docx(content: str, out_path: Path, today_str: str | None = None) -> None:
    today_str = today_str or datetime.now().strftime("%d.%m.%Y")
    content = _strip_md_escapes(content)

    # Split in Bloecke (Doppel-Newline = Absatz)
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    def add_paragraph(text: str, *, bold: bool = False, align=None) -> None:
        para = doc.add_paragraph()
        if align is not None:
            para.alignment = align
        run = para.add_run(text)
        if bold:
            run.bold = True

    def add_block(lines: list[str], *, align=None, bold: bool = False) -> None:
        """Mehrzeiliger Block (z.B. Adresse, Empfaenger): jede Zeile als
        eigener Absatz mit `space_after=0` damit's gestapelt aussieht."""
        for ln in lines:
            ln = ln.strip()
            if not ln:
                continue
            add_paragraph(ln, bold=bold, align=align)

    # Klassifizierung der Bloecke
    # Block 1: Absender (mehrzeilig) -> rechts
    # Block 2: Empfaenger
    # Block 3: Datum (Ulm, ...) -> rechts
    # Block 4: Betreff (mehrzeilig, bold)
    # Block 5: Anrede
    # Block 6+: Body (mit Section-Headers als bold-Einzeiler)

    if not paragraphs:
        raise ValueError("Kein Inhalt zum Konvertieren gefunden")

    blocks_used = 0

    # Block 1: Absender (rechts) — bis Block ohne "Universität Ulm"
    if blocks_used < len(paragraphs):
        add_block(paragraphs[blocks_used].split("\n"), align=WD_ALIGN_PARAGRAPH.RIGHT)
        blocks_used += 1
        add_paragraph("")

    # Block 2: Empfaenger (linksbuendig) — nur wenn passend
    if blocks_used < len(paragraphs) and (
        "International Office" in paragraphs[blocks_used]
        or "Universität" in paragraphs[blocks_used]
    ):
        add_block(paragraphs[blocks_used].split("\n"))
        blocks_used += 1
        add_paragraph("")

    # Block 3: Datum (rechts) — wenn das Block mit "Ulm," oder Datum aussieht
    if blocks_used < len(paragraphs) and re.match(
        r"^[A-ZÄÖÜ][a-zäöü]+,\s*\d{1,2}\.\d{1,2}\.\d{4}", paragraphs[blocks_used]
    ):
        add_paragraph(paragraphs[blocks_used], align=WD_ALIGN_PARAGRAPH.RIGHT)
        blocks_used += 1
        add_paragraph("")

    # Block 4: Betreff (bold, mehrzeilig)
    # Nehme den naechsten Block + ggf. einen weiteren wenn er auch ohne Punkt endet
    if blocks_used < len(paragraphs):
        subj = paragraphs[blocks_used]
        if any(w in subj for w in ("Motivationsschreiben", "Bewerbung", "Förderung", "Antrag")):
            add_block(subj.split("\n"), bold=True)
            blocks_used += 1
            add_paragraph("")

    # Rest: Body
    while blocks_used < len(paragraphs):
        block = paragraphs[blocks_used]
        next_blank = blocks_used + 1 < len(paragraphs)
        # Sektion-Header (einzeiliger Block, kurz, kein Punkt am Ende, gefolgt von Body)
        if "\n" not in block and _is_section_header(block, next_blank):
            add_paragraph("")
            add_paragraph(block, bold=True)
        else:
            add_paragraph(block)
        blocks_used += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def update_documents_json(doc_id: str) -> None:
    p = Path(r"C:/Users/Robin/Jarvis/Mission-Control/data/documents.json")
    if not p.exists():
        return
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return
    for d in data.get("documents", []):
        if d.get("id") == doc_id:
            d["last_synced"] = datetime.now().isoformat(timespec="seconds")
    tmp = p.with_suffix(p.suffix + ".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(p)


def main() -> int:
    if len(sys.argv) < 3:
        print("Usage: sync-doc-to-word.py <input-text-file> <output-docx> [--doc-id ID]")
        return 1
    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    doc_id = ""
    if "--doc-id" in sys.argv:
        i = sys.argv.index("--doc-id")
        doc_id = sys.argv[i + 1] if i + 1 < len(sys.argv) else ""

    content = in_path.read_text(encoding="utf-8")
    build_docx(content, out_path)
    print(f"OK -> {out_path} ({out_path.stat().st_size} bytes)")

    if doc_id:
        update_documents_json(doc_id)
        print(f"documents.json updated for {doc_id}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
