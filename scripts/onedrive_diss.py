"""
OneDrive Diss Word-Sync
Erstellt/aktualisiert Dissertation.docx in OneDrive via Microsoft Graph API.
Betreuer-Link wird zurückgegeben.
"""
import io
import json
import os
import pathlib
import sys

import msal
import requests

SCOPES = ["Files.ReadWrite", "Files.ReadWrite.All"]
GRAPH_BASE = "https://graph.microsoft.com/v1.0"
ONEDRIVE_PATH = "Dokumente/Studium/!!Medizin/03 Doktorarbeit/ITI/Dissertation-Draft.docx"
CACHE_FILE = pathlib.Path(__file__).parent / ".onedrive-token-cache.json"


def _load_cache():
    cache = msal.SerializableTokenCache()
    if CACHE_FILE.exists():
        cache.deserialize(CACHE_FILE.read_text(encoding="utf-8"))
    return cache


def _save_cache(cache):
    if cache.has_state_changed:
        CACHE_FILE.write_text(cache.serialize(), encoding="utf-8")


def _get_token(interactive=False):
    cache = _load_cache()
    client_id = os.getenv("MS_CLIENT_ID", "59386692-536e-4afe-9ae2-0f728d617727")
    app = msal.PublicClientApplication(
        client_id=client_id,
        authority="https://login.microsoftonline.com/common",
        token_cache=cache,
    )
    accounts = app.get_accounts()
    if accounts:
        result = app.acquire_token_silent(SCOPES, account=accounts[0])
        if result and "access_token" in result:
            _save_cache(cache)
            return result["access_token"]

    if not interactive:
        return None

    flow = app.initiate_device_flow(scopes=SCOPES)
    print(flow.get("message", "Login-Fehler"))
    result = app.acquire_token_by_device_flow(flow)
    _save_cache(cache)
    if "access_token" not in result:
        print("Login fehlgeschlagen:", result)
        sys.exit(1)
    return result["access_token"]


def _build_docx(content: str) -> bytes:
    """Erzeugt eine einfache .docx-Datei. Fallback auf Minimal-DOCX wenn python-docx fehlt."""
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading("Dissertation — ITI Ulm", 0)
        for block in content.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("# "):
                doc.add_heading(block[2:], level=1)
            elif block.startswith("## "):
                doc.add_heading(block[3:], level=2)
            elif block.startswith("### "):
                doc.add_heading(block[4:], level=3)
            else:
                p = doc.add_paragraph(block)
                p.style.font.size = Pt(11)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # Minimale DOCX-Struktur (ZIP) ohne externe Deps
        return _minimal_docx(content)


def _minimal_docx(content: str) -> bytes:
    import zipfile, textwrap
    body_xml = ""
    for line in content.splitlines():
        text = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        body_xml += f'<w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    doc_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:body>{body_xml}<w:sectPr/></w:body></w:document>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
    )
    ct = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", ct)
        zf.writestr("word/_rels/document.xml.rels", rels)
        zf.writestr("word/document.xml", doc_xml)
    return buf.getvalue()


def upload_to_onedrive(content: str, token: str) -> dict:
    docx_bytes = _build_docx(content)
    url = f"{GRAPH_BASE}/me/drive/root:/{ONEDRIVE_PATH}:/content"
    r = requests.put(
        url,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        data=docx_bytes,
        timeout=60,
    )
    if r.status_code not in (200, 201):
        return {"ok": False, "error": f"HTTP {r.status_code}: {r.text[:200]}"}
    item = r.json()
    item_id = item.get("id", "")

    # Sharing-Link erstellen (edit, damit Betreuer kommentieren können)
    share_url = f"{GRAPH_BASE}/me/drive/items/{item_id}/createLink"
    sr = requests.post(
        share_url,
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
        json={"type": "edit", "scope": "organization"},
        timeout=30,
    )
    share_link = ""
    if sr.status_code == 200:
        share_link = sr.json().get("link", {}).get("webUrl", "")

    web_url = item.get("webUrl", share_link)
    return {"ok": True, "webUrl": web_url, "shareLink": share_link, "itemId": item_id}


def get_status(token: str) -> dict:
    url = f"{GRAPH_BASE}/me/drive/root:/{ONEDRIVE_PATH}"
    r = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=20)
    if r.status_code == 404:
        return {"exists": False}
    if r.status_code != 200:
        return {"exists": False, "error": f"HTTP {r.status_code}"}
    item = r.json()
    return {
        "exists": True,
        "lastModified": item.get("lastModifiedDateTime", ""),
        "size": item.get("size", 0),
        "webUrl": item.get("webUrl", ""),
    }


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("cmd", choices=["login", "status", "upload"])
    parser.add_argument("--content", default="# Dissertation\n\nPlatzhalter.")
    args = parser.parse_args()

    if args.cmd == "login":
        tok = _get_token(interactive=True)
        print("Login erfolgreich.")
    elif args.cmd == "status":
        tok = _get_token()
        if not tok:
            print("Nicht eingeloggt. Bitte zuerst: python scripts/onedrive_diss.py login")
            sys.exit(1)
        print(json.dumps(get_status(tok), indent=2, ensure_ascii=False))
    elif args.cmd == "upload":
        tok = _get_token(interactive=True)
        result = upload_to_onedrive(args.content, tok)
        print(json.dumps(result, indent=2, ensure_ascii=False))
